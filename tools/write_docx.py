#!/usr/bin/env python3
"""MCP server that exposes a write_docx tool to Claude Code."""

import json
import sys
from docx import Document
from docx.shared import Pt
from mcp.server import Server
from mcp.server.stdio import stdio_server
from mcp import types

app = Server("docx-writer")

@app.list_tools()
async def list_tools():
    return [
        types.Tool(
            name="write_docx",
            description="Write a structured Word document with sections and subheadings",
            inputSchema={
                "type": "object",
                "properties": {
                    "filename": {"type": "string"},
                    "sections": {
                        "type": "array",
                        "items": {
                            "type": "object",
                            "properties": {
                                "heading": {"type": "string"},
                                "content": {"type": "string"},
                                "subheadings": {
                                    "type": "array",
                                    "items": {
                                        "type": "object",
                                        "properties": {
                                            "title": {"type": "string"},
                                            "content": {"type": "string"}
                                        }
                                    }
                                }
                            }
                        }
                    }
                },
                "required": ["filename", "sections"]
            }
        )
    ]

@app.call_tool()
async def call_tool(name: str, arguments: dict):
    if name == "write_docx":
        doc = Document()

        # Title
        doc.add_heading(arguments.get("title", "Newsletter Analysis"), 0)

        for section in arguments["sections"]:
            # H1
            doc.add_heading(section["heading"], level=1)

            if section.get("content"):
                doc.add_paragraph(section["content"])

            for sub in section.get("subheadings", []):
                # H2
                doc.add_heading(sub["title"], level=2)
                doc.add_paragraph(sub["content"])

            doc.add_paragraph()  # spacing

        filepath = f"output/{arguments['filename']}"
        doc.save(filepath)

        return [types.TextContent(
            type="text",
            text=f"Document saved to {filepath}"
        )]

if __name__ == "__main__":
    import asyncio
    asyncio.run(stdio_server(app))