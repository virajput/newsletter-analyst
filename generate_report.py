#!/usr/bin/env python3
"""Generate The Neural Blueprint growth analysis report as a Word document."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_heading(paragraph, level=1):
    """Apply color to heading paragraphs."""
    if level == 1:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(0xE8, 0x70, 0x3A)  # brand orange
    elif level == 2:
        for run in paragraph.runs:
            run.font.bold = True

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    set_heading(p, 1)
    return p

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    return p

def add_body(doc, text):
    return doc.add_paragraph(text)

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    return p

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light List Accent 1'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row_data in rows:
        row = table.add_row().cells
        for i, val in enumerate(row_data):
            row[i].text = val
    doc.add_paragraph()

# ── Title Page ───────────────────────────────────────────────────────────────
title = doc.add_heading('The Neural Blueprint', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0xE8, 0x70, 0x3A)

subtitle = doc.add_paragraph('Growth Analysis & Opportunity Report')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

date_p = doc.add_paragraph(f'Prepared: March 30, 2026  |  qubytes.substack.com')
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ── 1. Executive Summary ─────────────────────────────────────────────────────
add_h1(doc, '1. Executive Summary')
add_body(doc,
    'The Neural Blueprint (published at qubytes.substack.com) is a technically rigorous Substack newsletter '
    'targeting ML engineers, AI architects, and senior practitioners building production AI systems. With 500+ '
    'subscribers and weekly cadence, it occupies a defensible, high-value niche: research-grade AI content '
    'translated into actionable production guidance.')

add_body(doc,
    'Analysis of 10 top posts reveals a consistent editorial identity—problem-first, evidence-backed, '
    'production-oriented—delivered in a voice that reads as an insider sharing earned wisdom rather than a '
    'consultant selling frameworks. The newsletter is well-positioned for rapid growth and a paid tier, '
    'provided it doubles down on its core differentiation: depth that generic AI newsletters cannot match.')

add_h2(doc, 'Key Findings at a Glance')
bullets = [
    'Strongest content: multi-agent protocol deep-dives and research paper breakdowns for practitioners',
    'Top engagement drivers: numbered frameworks, production failure stories, and honest tradeoff analysis',
    'Critical gap: no owned SEO asset (glossary, tooling reference) to capture search traffic',
    'Format gap: long-form dominates; short-form and video repurposing nearly absent',
    'Paid tier opportunity: checklists, decision frameworks, and implementation guides are under-gated',
    'Audience ready for segmentation: beginner practitioners vs. senior architects have different needs',
]
for b in bullets:
    add_bullet(doc, b)

doc.add_paragraph()

# ── 2. Audience Profile ──────────────────────────────────────────────────────
add_h1(doc, '2. Audience Profile')
add_body(doc,
    'Based on content signals, writing assumptions, and CTA language across all 10 posts, the core '
    'readership is:')

add_h2(doc, 'Primary Persona: The Production AI Architect')
add_bullet(doc, 'Role: ML engineer, AI architect, technical lead, or senior software engineer')
add_bullet(doc, 'Experience: 5–15 years in software; 1–4 years working with LLMs/agents')
add_bullet(doc, 'Context: Working in mid-to-large engineering teams; often at companies beginning AI adoption')
add_bullet(doc, 'Challenge: Translating AI research and demos into reliable, cost-effective production systems')
add_bullet(doc, 'Frustration: Most AI content is either too academic (papers) or too shallow (blog posts)')

add_h2(doc, 'Secondary Persona: The Technical Decision-Maker')
add_bullet(doc, 'Role: CTO, VP Engineering, Principal Architect evaluating AI investments')
add_bullet(doc, 'Need: Defensible technical reasoning for tool selection and architecture decisions')
add_bullet(doc, 'Pain: Vendor marketing obscures real tradeoffs; analyst reports lack engineering depth')

add_h2(doc, 'Core Audience Frustrations')
add_table(doc,
    ['Frustration', 'Evidence in Posts'],
    [
        ['AI demos that don\'t survive production', '"The architecture that survives production" (tagline); doom-loop detection; instruction fade-out discussion'],
        ['Cost overruns from runaway agents', '$47K runaway loop story in long-horizon agents post'],
        ['Tool selection paralysis', 'Long-horizon agent comparison; 5-tool productivity system'],
        ['Inability to keep up with research', 'Agent Lightning, Gnosis, LLaDA paper breakdowns'],
        ['Black-box AI opacity and trust problems', 'AG-UI transparency section; Gnosis error detection post'],
        ['Architecture decisions that create future debt', 'Scalable systems 7-layer framework; tight coupling discussion'],
    ]
)

doc.add_paragraph()

# ── 3. Problems Solved ───────────────────────────────────────────────────────
add_h1(doc, '3. Problems Solved')
add_body(doc, 'Five primary problem categories surface consistently across top posts:')

problems = [
    ('P1: How do I design multi-agent systems that don\'t fail in production?',
     '"Understanding Multi-Agent Communication" (3-protocol stack); "The Hidden Engineering Behind AI Coding Agents" (doom loops, instruction fade-out, 5-layer safety). The newsletter provides architecture diagrams, failure mode catalogs, and decision frameworks.'),
    ('P2: How do I evaluate and choose between competing AI tools/frameworks?',
     '"The Long-Horizon Agent Race" (4-product comparison with security posture); "I Replaced ChatGPT with 5 AI Tools" (specialization-over-generalization argument). Posts include role-based selection matrices and real incident examples.'),
    ('P3: How do I stay current with AI research without reading 50 papers a month?',
     '"Can AI Catch Its Own Mistakes?" (Gnosis); "Agent Lightning" (LightningRL framework); "LLaDA2.1" (diffusion LM speed). Research papers distilled into practitioner-relevant summaries with performance benchmarks and applicability assessments.'),
    ('P4: How do I build systems that scale—both AI and traditional?',
     '"Scalable Systems Guide: 7 Layers Every Dev Should Know" (Netflix case study); "Agentic Business Process Management Systems" (5-layer pyramid). Frameworks that map complexity into independently scalable layers.'),
    ('P5: What do production AI governance and safety architectures look like?',
     '"Beyond Core Protocols" (AP2 authorization, spending governance); "Hidden Engineering" (5-layer safety architecture). Practical patterns with code examples, not policy documents.'),
]

for title_text, evidence in problems:
    add_h2(doc, title_text)
    add_body(doc, evidence)

doc.add_paragraph()

# ── 4. Content Themes & Patterns ────────────────────────────────────────────
add_h1(doc, '4. Content Themes & Patterns')

add_h2(doc, 'Theme Distribution Across Top 10 Posts')
add_table(doc,
    ['Theme', 'Posts', 'Avg. Depth'],
    [
        ['Multi-agent protocols & architecture', '2 posts (MCP/A2A/AG-UI series)', 'Very deep (2,800–3,500 words)'],
        ['Research paper breakdowns', '3 posts (Agent Lightning, Gnosis, LLaDA)', 'Medium (900–2,000 words)'],
        ['Tool selection & comparison', '2 posts (long-horizon agents, AI tools)', 'Deep (2,400 words)'],
        ['Traditional systems architecture', '2 posts (7-layer scalability, Diagram-as-Code)', 'Medium (1,800 words)'],
        ['Process automation & BPM', '1 post (agentic BPM)', 'Light (~620 words)'],
    ]
)

add_h2(doc, 'Recurring Structural Patterns')
patterns = [
    'Problem-first framing: Every post opens with a production failure, frustration, or gap—never a solution',
    'Named frameworks: Top posts introduce a named model (7-Layer Model, 3-Protocol Stack, 5-Layer Safety) that readers can reference and share',
    'Limitations transparency: Each post explicitly states what the solution does NOT handle—builds trust',
    'War stories: Concrete failure incidents ($47K loop, 47% database deletion, 3am pages) establish credibility',
    'Tradeoff tables: Visual comparison of old vs. new approaches with specific dimensions',
    'Research-to-practice bridge: Academic papers translated into "what this means for your architecture"',
    'Actionable endings: Posts close with next steps, decision frameworks, or implementation checklists',
]
for p in patterns:
    add_bullet(doc, p)

doc.add_paragraph()

# ── 5. Writing Style Analysis ────────────────────────────────────────────────
add_h1(doc, '5. Writing Style Analysis')

add_h2(doc, 'Voice & Tone DNA')
add_body(doc,
    'The Neural Blueprint\'s voice is that of a senior AI architect sharing battlefield knowledge—not an '
    'educator lecturing or a journalist reporting. Key attributes:')

voice_items = [
    ('Earned authority', 'Speaks from first-person experience with specific systems and failures, not from credentials'),
    ('Production bias', 'Every concept is filtered through "does this survive real load, real failures, real constraints?"'),
    ('No-hype clarity', 'Explicitly refuses vendor marketing: "No vendor marketing. No protocol wars. Just clear architectural guidance."'),
    ('Honest limitations', 'Freely acknowledges what frameworks don\'t solve, creating higher trust for what they do solve'),
    ('Accessible depth', 'Complex concepts explained via analogy without dumbing down—recipes vs. AI agents, vital signs vs. error detection'),
]

for attr, desc in voice_items:
    p = doc.add_paragraph()
    run = p.add_run(f'{attr}: ')
    run.bold = True
    p.add_run(desc)

add_h2(doc, 'Post Anatomy (Typical Structure)')
structure_steps = [
    '1. Hook: Production failure story or counterintuitive claim (2–4 sentences)',
    '2. Problem statement: What\'s broken and why existing approaches fail',
    '3. Framework introduction: Named model with 3–7 components',
    '4. Component breakdown: Each part with use case, mechanics, and limitations',
    '5. Integration example: Realistic scenario showing all components working together',
    '6. Tradeoff acknowledgment: What the framework doesn\'t solve',
    '7. Actionable close: Decision framework, checklist, or "this week\'s homework"',
]
for step in structure_steps:
    add_bullet(doc, step)

add_h2(doc, 'Word Count Distribution')
add_table(doc,
    ['Post Type', 'Typical Word Count', 'Examples'],
    [
        ['Deep architecture breakdown', '2,800–3,500 words', 'Multi-agent protocols, hidden AI engineering'],
        ['Research paper breakdown', '900–2,000 words', 'Gnosis, Agent Lightning'],
        ['Tool comparison/productivity', '2,400–2,600 words', 'Long-horizon agents, AI tools stack'],
        ['Light explainer', '600–800 words', 'Agentic BPM'],
    ]
)

doc.add_paragraph()

# ── 6. Content Gaps & Opportunities ─────────────────────────────────────────
add_h1(doc, '6. Content Gaps & Opportunities')

add_h2(doc, 'Format Gaps (Underutilized Given What Resonates)')
format_gaps = [
    ('Weekly research digest', 'Three research breakdowns already among top posts; a dedicated weekly series would compound SEO and establish a recurring appointment with readers'),
    ('Interactive reference tools', 'Decision frameworks appear in posts as text; a standalone "Which agent framework?" quiz or calculator would drive recurring visits'),
    ('Short-form explainers (< 800 words)', 'Agentic BPM at 620 words got engagement; quick takes on emerging concepts could run 2–3x weekly cadence'),
    ('Annotated code repositories', 'Implementation-focused posts reference GitHub but don\'t link to maintained companion repos—a gap for paid tier'),
    ('Visual-first content', 'Architecture diagrams are mentioned but not shown; rich diagram posts would perform strongly on LinkedIn'),
]
for gap, rationale in format_gaps:
    p = doc.add_paragraph()
    run = p.add_run(f'{gap}: ')
    run.bold = True
    p.add_run(rationale)

add_h2(doc, 'Content Gaps (Missing Angles the Audience Would Clearly Value)')
content_gaps = [
    'Agent testing and evaluation: How do you test a system that is non-deterministic? No coverage yet.',
    'Prompt engineering for production: Not the "write better prompts" kind—the system prompt architecture, injection defense, and fade-out mitigation kind.',
    'Cost modeling and FinOps for AI: $47K runaway loops are mentioned but no post systematically addresses budgeting and cost prediction.',
    'Real enterprise case studies: Hypothetical scenarios dominate; anonymized real deployments would be far more valuable.',
    'Data engineering for RAG: High demand topic; complementary to existing agent architecture content.',
    'Security and red-teaming: Supply chain risks mentioned once; a dedicated security architecture post is missing.',
    'Agent monitoring and observability: AG-UI covers events conceptually but no post covers observability stacks (OpenTelemetry, tracing).',
]
for gap in content_gaps:
    add_bullet(doc, gap)

doc.add_paragraph()

# ── 7. Ten Specific Content Ideas ────────────────────────────────────────────
add_h1(doc, '7. Ten Specific Content Ideas')
add_body(doc,
    'Each idea passes two tests: (1) maps to a pattern observed in top posts, and (2) directly addresses '
    'an audience frustration identified in Phase 2.')

ideas = [
    {
        'title': 'The Production Checklist for Multi-Agent Systems: 15 Things to Verify Before Going Live',
        'why': 'Maps to the multi-agent protocols post (doom loops, instruction fade-out, delegation failures) and the hidden engineering post (5-layer safety). Addresses the #1 frustration: AI demos that fail in production.',
        'format': 'Checklist / breakdown',
        'hook': 'Your multi-agent system works in the demo environment. Here are the 15 failure modes that will humiliate it in production.',
    },
    {
        'title': 'MCP vs. LangChain vs. AutoGen: The Architecture Decision Matrix for Agent Orchestration',
        'why': 'Maps to Agent Lightning post (all three frameworks tested head-to-head) and multi-agent protocols post. Addresses tool selection paralysis—audiences facing this decision Google exactly this comparison.',
        'format': 'Comparison with decision matrix',
        'hook': 'Three senior engineers gave me three different answers about which framework to use. Here is the decision matrix that ends the debate.',
    },
    {
        'title': 'How to Calculate Real AI Agent Costs Before They Bankrupt You',
        'why': 'Maps to long-horizon agent race ($47K runaway loop) and AI tools post ($150/month, 82 hours saved). Addresses the cost overrun frustration with a practical formula and worked examples.',
        'format': 'Tutorial with cost model template',
        'hook': 'A $47,000 runaway agent loop is not a bug. It is the cost of not building a budget model before deployment.',
    },
    {
        'title': 'The Context Window Crisis: How Production Agents Actually Remember Across 100-Turn Sessions',
        'why': 'Maps directly to hidden engineering post (ACC pipeline, 54% context reduction, instruction fade-out after 15 tool calls). Addresses the gap between demo agent behavior and production agent behavior.',
        'format': 'Technical deep-dive with diagrams',
        'hook': 'After 15 tool calls, your agent starts forgetting its own safety rules. The fix is not a longer context window.',
    },
    {
        'title': 'I Read 40 AI Papers This Month: The 5 That Will Change How You Build in 2026',
        'why': 'Maps to the three research breakdown posts (Gnosis, Agent Lightning, LLaDA). Addresses inability to keep up with research. A monthly digest format creates a recurring appointment and compound SEO value.',
        'format': 'Research digest / roundup',
        'hook': 'I filtered 40 papers down to five that have direct architectural implications for production systems. Here is what you missed.',
    },
    {
        'title': 'Enterprise AI Governance: The 8 Questions Your Legal Team Will Ask Before Deploying Agents',
        'why': 'Maps to long-horizon agents post (regulated industries, database deletion incidents) and agentic commerce post (spending governance, cryptographic authorization). Addresses the governance gap for enterprise readers.',
        'format': 'Framework / checklist',
        'hook': 'Your legal team will not block your AI deployment because the model is bad. They will block it because you cannot answer these eight questions.',
    },
    {
        'title': 'Building an Observable Agent: The 12 AG-UI Events Every Production System Should Emit',
        'why': 'Maps directly to the multi-agent communications post (AG-UI event taxonomy: TASK_STARTED, DECISION_POINT, HUMAN_INPUT_NEEDED, etc.). Addresses AI opacity and trust—turns conceptual coverage into a practical reference.',
        'format': 'Reference guide with code examples',
        'hook': 'If your agent cannot explain what it is doing while it is doing it, you do not have a production system. You have a black box with consequences.',
    },
    {
        'title': 'From Reactive to Autonomous: A 4-Stage Agent Maturity Model for Enterprise Teams',
        'why': 'Maps to long-horizon agents post (reactive vs. autonomous execution model) and agentic BPM post (triage patterns, supervision models). Addresses decision-makers evaluating how far to push agent autonomy.',
        'format': 'Framework / maturity model',
        'hook': 'Most enterprise AI implementations are stuck at Stage 1. Here is what Stages 2, 3, and 4 look like—and the specific engineering gates between them.',
    },
    {
        'title': 'The Open-Source Agent Security Audit: What to Check Before Trusting an MIT-Licensed Agent With Your Terminal',
        'why': 'Maps to long-horizon agents post (OpenClaw vulnerabilities, PyPI malware fork bomb). Addresses security-conscious readers at companies that cannot use closed-source agents.',
        'format': 'Security checklist / case study',
        'hook': 'A malicious PyPI package caused a fork bomb on an engineer\'s machine last quarter. Your open-source agent has the same attack surface.',
    },
    {
        'title': 'The Specialized AI Stack: How to Get 3x Productivity Without Paying for Every Premium Tier',
        'why': 'Maps to AI tools productivity post (5-tool system, 82 hours/month saved, tiered adoption path). Addresses the accessibility concern—many readers want results at lower cost.',
        'format': 'Tutorial with tiered options',
        'hook': 'The $150/month full-stack approach is not the starting point. Here is the $20/month version that gets you 80% of the results.',
    },
]

for i, idea in enumerate(ideas, 1):
    add_h2(doc, f'Idea {i}: {idea["title"]}')
    p = doc.add_paragraph()
    p.add_run('Why it works: ').bold = True
    p.add_run(idea['why'])
    p = doc.add_paragraph()
    p.add_run('Format: ').bold = True
    p.add_run(idea['format'])
    p = doc.add_paragraph()
    p.add_run('Hook sentence: ').bold = True
    p.add_run(f'"{idea["hook"]}"')
    doc.add_paragraph()

# ── 8. Growth Strategy Recommendations ───────────────────────────────────────
add_h1(doc, '8. Growth Strategy Recommendations')

add_h2(doc, '8.1 Newsletter-Level SEO')
add_body(doc,
    'SEO strategy should operate at the newsletter level—building topical authority in a cluster, '
    'not chasing individual post keywords.')

add_body(doc, 'Niche Territory: Production AI system architecture for engineers.')

add_h2(doc, 'Priority Keyword Clusters')
add_table(doc,
    ['Cluster', 'Target Terms', 'Rationale'],
    [
        ['Multi-agent architecture', '"multi-agent system architecture," "agent orchestration production," "MCP A2A protocol"', 'Two top posts already cover this; build topical depth'],
        ['LLM production engineering', '"LLM production deployment," "AI agent failure modes," "context window management production"', 'Underserved by most AI content; high engineer search intent'],
        ['AI research for practitioners', '"AI paper explained," "ML research breakdown," "agent training RL"', 'High volume; compound SEO with weekly research series'],
        ['Enterprise AI governance', '"enterprise AI agent security," "autonomous agent governance," "AI spending controls"', 'Low competition; high commercial intent from enterprise readers'],
        ['AI cost optimization', '"AI agent cost model," "LLM inference cost," "runaway agent cost prevention"', 'Growing concern as agent deployments scale; practical post angle'],
    ]
)

add_h2(doc, 'Owned Asset Recommendation: The Production AI Architecture Glossary')
add_body(doc,
    'Create a standalone, searchable web page (hosted on Substack or a linked subdomain) containing '
    'definitions for every protocol, pattern, and concept covered in the newsletter: MCP, A2A, AG-UI, '
    'ACC, LightningRL, doom loop detection, instruction fade-out, etc.')
add_bullet(doc, 'SEO value: Engineers Googling "what is AG-UI protocol" land on a Neural Blueprint page')
add_bullet(doc, 'Retention value: Subscribers bookmark it as a reference; it creates recurring visits')
add_bullet(doc, 'Positioning value: Owning the definitions of emerging terminology establishes editorial authority')
add_bullet(doc, 'Effort: Each existing post already contains definitions—this is a curation and formatting task')

add_h2(doc, '8.2 Content Repurposing Strategy')
add_table(doc,
    ['Post Type', 'Best Platform', 'Format', 'Example'],
    [
        ['Multi-agent protocol posts', 'LinkedIn', 'Protocol stack carousel (10 slides)', 'MCP/A2A/AG-UI three-layer diagram'],
        ['Research paper breakdowns', 'Twitter/X', 'Thread (8–10 tweets, one finding per tweet)', '"Gnosis: a 5M parameter model that outperforms billion-parameter checkers"'],
        ['Tool comparison posts', 'LinkedIn + Twitter', 'Table screenshot + "the one I\'d pick" take', 'Long-horizon agent comparison matrix'],
        ['Architecture breakdowns', 'YouTube Shorts', '60-second "the problem / the fix" format', '"Why agents forget their safety rules after 15 calls"'],
        ['Productivity frameworks', 'LinkedIn', 'Personal story post with framework reveal', '5-tool AI stack origin story'],
    ]
)

add_h2(doc, '8.3 Audience Segmentation Strategy')
add_body(doc,
    'The audience naturally splits into two segments with different needs and different willingness to pay:')

add_h2(doc, 'Segment A: The Practitioner (0–3 years AI experience)')
add_bullet(doc, 'Needs: Conceptual clarity, tool introductions, research translated to plain English')
add_bullet(doc, 'Content: Research digests, "what is X" explainers, productivity frameworks')
add_bullet(doc, 'Channel: Free tier, Twitter/X, LinkedIn')

add_h2(doc, 'Segment B: The Architect (3+ years, building production systems)')
add_bullet(doc, 'Needs: Implementation patterns, failure mode analysis, governance frameworks, production checklists')
add_bullet(doc, 'Content: Deep architecture breakdowns, annotated code repos, decision matrices')
add_bullet(doc, 'Channel: Paid tier candidate; highest subscriber lifetime value')

add_body(doc,
    'Recommended approach: Keep Segment A content free to maximize top-of-funnel. Gate Segment B content '
    '(checklists, implementation guides, companion code repos) behind a paid tier priced at $9–$12/month or $90/year.')

doc.add_paragraph()

# ── 9. Competitive Positioning ───────────────────────────────────────────────
add_h1(doc, '9. Competitive Positioning')

add_h2(doc, 'Competitive Landscape')
add_table(doc,
    ['Newsletter / Resource', 'Positioning', 'The Neural Blueprint Advantage'],
    [
        ['The Batch (deeplearning.ai)', 'Research news digest, broad AI coverage', 'Deeper architecture focus; production-grade not research-grade'],
        ['Ahead of AI', 'Academic paper summaries by Sebastian Raschka', 'More production-focused; less academic tone'],
        ['Latent Space', 'AI researcher interviews, podcast-first', 'Written-first; architecture diagrams; practitioner angle'],
        ['TLDR AI', 'Short daily news digest', 'Depth over breadth; searchable reference content'],
        ['Interconnects', 'Policy + research commentary', 'Engineering-first; implementation patterns not policy analysis'],
    ]
)

add_h2(doc, 'Defensible Differentiation')
add_body(doc,
    'The Neural Blueprint\'s most defensible position is the intersection of three attributes that no '
    'competitor currently occupies simultaneously:')
add_bullet(doc, 'Production-focused (not demo or research-focused)')
add_bullet(doc, 'Architecture-depth (not news summaries or surface-level takes)')
add_bullet(doc, 'Emerging protocols (MCP, A2A, AG-UI, ACP—before they become mainstream)')

add_body(doc,
    'This creates a 12–18 month window where the newsletter can establish itself as the definitive '
    'resource for production multi-agent architecture before the space becomes crowded. The tagline—'
    '"Built for engineers designing real AI systems. Not theory. Not demos. The architecture that '
    'survives production."—should be reinforced in every piece of content and every distribution channel.')

doc.add_paragraph()

# ── 10. Top 5 Next Steps ─────────────────────────────────────────────────────
add_h1(doc, '10. Top 5 Next Steps')
add_body(doc, 'Priority ordered. Each step has a one-sentence rationale and a concrete first action.')

next_steps = [
    (
        '1. Launch a weekly research digest series (starting within 2 weeks)',
        'Research paper breakdowns are already among the most-engaged posts; a named weekly series (e.g., "This Week in AI Architecture") compounds SEO, creates a subscriber appointment, and is low-lift to produce.',
        'First action: Publish the first edition covering 3–5 papers from the past 2 weeks with the same breakdown format as the Gnosis and Agent Lightning posts.',
    ),
    (
        '2. Build the Production AI Architecture Glossary (within 30 days)',
        'An owned, searchable reference page for protocols and patterns (MCP, A2A, AG-UI, instruction fade-out, doom loop detection) would capture significant search traffic and create recurring visits—content already exists, it just needs to be curated.',
        'First action: Create a Substack page (or linked microsite) with definitions extracted from the top 5 posts; share it in the next newsletter as a new subscriber resource.',
    ),
    (
        '3. Launch LinkedIn content strategy using existing posts (within 1 week)',
        'The multi-agent protocol posts contain the exact kind of architecture diagrams and frameworks that perform exceptionally well as LinkedIn carousels; this is the highest-distribution channel for the target audience with zero new content creation required.',
        'First action: Convert the 3-protocol stack diagram from the multi-agent post into a 10-slide LinkedIn carousel; post it and measure impressions vs. newsletter post engagement.',
    ),
    (
        '4. Design and announce the paid tier structure (within 60 days)',
        'The newsletter already produces paid-tier-quality content (implementation guides, decision frameworks, production checklists); formalizing access tiers with gated companion code repos and extended breakdowns would capture value from the highest-intent Segment B readers.',
        'First action: Survey current subscribers with 3 questions about what they would pay for; announce founding member pricing at a 40% discount for the first 100 subscribers.',
    ),
    (
        '5. Publish the Production Checklist for Multi-Agent Systems (within 2 weeks)',
        'This content idea has the highest immediate utility for the core audience, directly addresses the #1 reader frustration, and synthesizes material already covered across 3 top posts—making it low research lift and high distribution potential.',
        'First action: Draft the 15-point checklist pulling directly from the multi-agent protocols, hidden engineering, and long-horizon agents posts; format it as a gated PDF download to build the email list.',
    ),
]

for title_text, rationale, action in next_steps:
    add_h2(doc, title_text)
    p = doc.add_paragraph()
    p.add_run('Rationale: ').bold = True
    p.add_run(rationale)
    p = doc.add_paragraph()
    p.add_run('First action: ').bold = True
    p.add_run(action)
    doc.add_paragraph()

# ── Save ─────────────────────────────────────────────────────────────────────
filepath = 'output/neural_blueprint_growth_analysis.docx'
doc.save(filepath)
print(f'Document saved to {filepath}')
